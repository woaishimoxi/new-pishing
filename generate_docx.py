"""
生成论文大纲DOCX文档
"""
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

def create_thesis_outline():
    doc = Document()
    
    # 设置默认字体
    doc.styles['Normal'].font.name = 'SimSun'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    doc.styles['Normal'].font.size = Pt(12)
    
    # 标题样式
    def add_title(text, level=1):
        if level == 0:
            p = doc.add_heading(text, 0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'SimHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
                run.font.size = Pt(22)
        elif level == 1:
            p = doc.add_heading(text, 1)
            for run in p.runs:
                run.font.name = 'SimHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
                run.font.size = Pt(16)
        elif level == 2:
            p = doc.add_heading(text, 2)
            for run in p.runs:
                run.font.name = 'SimHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
                run.font.size = Pt(14)
        elif level == 3:
            p = doc.add_heading(text, 3)
            for run in p.runs:
                run.font.name = 'SimHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
                run.font.size = Pt(12)
        return p
    
    def add_para(text, indent=False):
        p = doc.add_paragraph(text)
        if indent:
            p.paragraph_format.first_line_indent = Cm(0.74)
        return p
    
    # 主标题
    add_title('面向中小企业的轻量化钓鱼邮件检测与溯源系统', 0)
    add_title('毕业论文大纲（修订版）', 1)
    
    # 第一章
    add_title('第一章 绪论', 1)
    
    add_title('1.1 研究背景与意义', 2)
    add_title('1.1.1 中小企业面临的网络安全现状', 3)
    add_para('• 钓鱼攻击占网络攻击比例超过90%（引用APWG、Verizon DBIR数据）', True)
    add_para('• 中小企业因资源有限成为攻击主要目标', True)
    
    add_title('1.1.2 传统安全设备在中小企业落地的痛点', 3)
    add_para('• 硬件成本高：传统邮件网关设备动辄数万元', True)
    add_para('• 运维门槛高：需要专业安全团队', True)
    add_para('• 资源消耗大：不适合中小企业轻量级IT环境', True)
    
    add_title('1.1.3 研究轻量化检测系统的必要性', 3)
    add_para('• 低成本部署需求', True)
    add_para('• 高检测精度要求', True)
    add_para('• 检测与溯源一体化需求', True)
    
    add_title('1.2 国内外研究现状', 2)
    add_title('1.2.1 邮件检测技术综述', 3)
    add_para('• 规则匹配方法：SpamAssassin、Rspamd', True)
    add_para('• 机器学习方法：朴素贝叶斯、SVM、随机森林', True)
    add_para('• 深度学习方法：TextCNN、BERT、GPT', True)
    
    add_title('1.2.2 邮件溯源技术发展现状', 3)
    add_para('• 邮件头分析技术', True)
    add_para('• IP地理位置追踪', True)
    add_para('• DNS黑名单检测（DNSBL）', True)
    
    add_title('1.2.3 现有研究的不足', 3)
    add_para('• 重检测轻溯源：多数系统缺少完整溯源能力', True)
    add_para('• 模型过于庞大：BERT等模型不适合中小企业部署', True)
    add_para('• 缺乏人机协同：无法处理边界样本', True)
    
    add_title('1.3 论文主要工作与创新点', 2)
    add_para('创新点1：提出了"LightGBM模型 + 规则引擎 + URL分析"的三维度融合检测架构', True)
    add_para('  - 模型轻量化：参数量小，推理速度快', True)
    add_para('  - 多维度融合：模型预测(60%) + URL风险(40%)', True)
    add_para('  - 智能降权：白名单URL、邮件认证通过自动降权', True)
    add_para('创新点2：设计了Kill Switch硬性规则机制', True)
    add_para('  - 沙箱检测恶意代码直接判定', True)
    add_para('  - 三重认证失败+身份冒充直接判定', True)
    add_para('  - VirusTotal高检测率直接判定', True)
    add_para('创新点3：实现了检测与溯源的闭环系统', True)
    add_para('  - IP地理位置追踪', True)
    add_para('  - 多DNSBL并行检测', True)
    add_para('  - 攻击链可视化展示', True)
    
    add_title('1.4 论文结构安排', 2)
    
    # 第二章
    add_title('第二章 相关技术与理论基础', 1)
    
    add_title('2.1 钓鱼邮件攻击原理与常见手法', 2)
    add_title('2.1.1 伪装技术（Spoofing）与社会工程学诱导', 3)
    add_para('• 发件人伪造', True)
    add_para('• 显示名称欺骗', True)
    add_para('• 相似域名欺骗', True)
    
    add_title('2.1.2 恶意链接与URL重定向机制', 3)
    add_para('• IP地址URL', True)
    add_para('• 短链接服务', True)
    add_para('• 品牌仿冒域名', True)
    
    add_title('2.2 邮件安全协议分析', 2)
    add_title('2.2.1 SPF、DKIM、DMARC协议原理', 3)
    add_para('• SPF：发件人IP验证', True)
    add_para('• DKIM：邮件签名验证', True)
    add_para('• DMARC：策略聚合与报告', True)
    
    add_title('2.2.2 MIME协议与邮件结构解析', 3)
    add_para('• 多部分邮件结构', True)
    add_para('• 附件编码方式', True)
    add_para('• HTML邮件解析', True)
    
    add_title('2.3 轻量化机器学习模型', 2)
    add_title('2.3.1 LightGBM梯度提升框架', 3)
    add_para('• 直方图加速算法', True)
    add_para('• 叶子生长策略', True)
    add_para('• 模型压缩与量化', True)
    
    add_title('2.3.2 特征工程与降维技术', 3)
    add_para('• 特征选择方法', True)
    add_para('• 特征重要性分析', True)
    
    add_title('2.4 威胁情报与溯源技术', 2)
    add_para('• 2.4.1 VirusTotal多引擎检测', True)
    add_para('• 2.4.2 DNS黑名单（DNSBL）查询', True)
    add_para('• 2.4.3 IP地理位置查询API', True)
    
    # 第三章
    add_title('第三章 系统需求分析与总体设计', 1)
    
    add_title('3.1 系统需求分析（面向中小企业场景）', 2)
    add_title('3.1.1 功能性需求', 3)
    add_para('• 邮件检测：支持手动输入、文件上传、IMAP自动拉取', True)
    add_para('• 溯源分析：IP追踪、黑名单检测、攻击链构建', True)
    add_para('• 域名管理：黑白名单配置、IOC规则管理', True)
    add_para('• 对抗测试：多种攻击类型模拟、鲁棒性评估', True)
    add_para('• 数据可视化：统计大屏、检测报告', True)
    
    add_title('3.1.2 非功能性需求', 3)
    add_para('• 低资源占用：内存<500MB，无需GPU', True)
    add_para('• 高实时性：单封邮件检测<2秒', True)
    add_para('• 易用性：Web界面，零配置启动', True)
    add_para('• 可扩展性：模块化设计，支持API扩展', True)
    
    add_title('3.2 系统总体架构设计', 2)
    add_title('3.2.1 系统架构图', 3)
    add_para('展示层（前端页面）→ API路由层（Flask Blueprint）→ 业务服务层（检测引擎、特征提取、溯源服务）→ 数据与模型层（LightGBM、SQLite、配置文件）', True)
    
    add_title('3.2.2 系统工作流程图', 3)
    add_para('邮件输入 → 邮件解析 → 特征提取 → 检测引擎 → 溯源分析 → 结果存储', True)
    
    add_title('3.3 系统模块划分', 2)
    add_para('• 3.3.1 邮件采集模块：手动输入、文件上传、IMAP协议自动拉取', True)
    add_para('• 3.3.2 邮件解析模块：邮件头解析、正文提取、附件解析、URL提取', True)
    add_para('• 3.3.3 特征提取模块：邮件头特征(8维)、URL特征(12维)、文本特征(6维)、附件特征(8维)、HTML特征(5维)', True)
    add_para('• 3.3.4 检测引擎模块：LightGBM模型预测、规则评分引擎、三维度融合决策', True)
    add_para('• 3.3.5 溯源分析模块：IP地理位置追踪、DNSBL黑名单检测、攻击链构建', True)
    add_para('• 3.3.6 前端展示模块：数据大屏、控制台、配置管理', True)
    
    # 第四章
    add_title('第四章 关键模块详细设计与实现（重点章节）', 1)
    
    add_title('4.1 邮件解析模块实现', 2)
    add_para('• 4.1.1 邮件结构解析：基于Python email库实现、多部分邮件处理、字符编码自动检测', True)
    add_para('• 4.1.2 发件人信息提取：显示名称与邮箱分离、域名提取与验证', True)
    add_para('• 4.1.3 URL与附件提取：正则表达式URL提取、HTML链接提取器、附件风险类型识别', True)
    
    add_title('4.2 特征工程实现', 2)
    add_title('4.2.1 邮件头特征（8维）', 3)
    add_para('is_suspicious_from_domain（可疑发件人域名）、received_hops_count（邮件跳转次数）、first_external_ip_is_blacklisted（源IP黑名单标志）、spf_fail（SPF认证失败）、dkim_fail（DKIM认证失败）、dmarc_fail（DMARC认证失败）、from_display_name_mismatch（显示名称不匹配）、from_domain_in_subject（域名出现在主题）', True)
    
    add_title('4.2.2 URL特征（12维）', 3)
    add_para('domain_age_days（域名年龄）、has_https（HTTPS标志）、is_short_url（短链接标志）、vt_detection_ratio（VirusTotal检测率）、is_ip_address（IP地址URL）、has_port（非标准端口）、url_length（URL长度）、has_suspicious_params（可疑参数）、has_at_symbol（@符号）、has_subdomain（子域名）、path_depth（路径深度）、query_length（查询长度）', True)
    
    add_title('4.2.3 文本特征（6维）', 3)
    add_para('紧急关键词计数、金融关键词计数、文本长度、感叹号数量、大写字母比例、紧急度评分', True)
    
    add_title('4.2.4 附件特征（8维）', 3)
    add_para('附件数量、高危扩展名标志、双重扩展名标志、可执行文件标志、文件大小、MIME类型等', True)
    
    add_title('4.2.5 HTML特征（5维）', 3)
    add_para('HTML标签数量、链接数量、表单数量、隐藏链接标志、外部资源引用', True)
    
    add_title('4.3 检测引擎实现', 2)
    add_title('4.3.1 LightGBM模型训练', 3)
    add_para('• 数据集：Enron + SpamAssassin + 自建钓鱼邮件集', True)
    add_para('• 特征标准化处理', True)
    add_para('• 模型参数调优', True)
    add_para('• 模型文件大小：<1MB', True)
    
    add_title('4.3.2 Kill Switch硬性规则', 3)
    add_para('规则1: 沙箱检测到恶意代码 → 返回PHISHING', True)
    add_para('规则2: 三重认证失败 + 身份冒充 → 返回PHISHING', True)
    add_para('规则3: VirusTotal高检测率(>50%) → 返回PHISHING', True)
    
    add_title('4.3.3 三维度融合算法', 3)
    add_para('• 模型预测（权重60%）+ URL风险（权重40%）', True)
    add_para('• 白名单URL降权：-0.5', True)
    add_para('• 邮件认证通过降权：-0.3', True)
    
    add_title('4.3.4 阈值判定', 3)
    add_para('• 置信度 ≥ 0.70：PHISHING（钓鱼）', True)
    add_para('• 置信度 ≥ 0.40：SUSPICIOUS（可疑）', True)
    add_para('• 置信度 < 0.40：SAFE（安全）', True)
    
    add_title('4.4 溯源分析实现', 2)
    add_para('• 4.4.1 IP地理位置追踪：从Received头提取源IP、调用IP-API获取地理位置', True)
    add_para('• 4.4.2 DNSBL黑名单检测：支持5+主流DNSBL服务、ThreadPoolExecutor并行查询', True)
    add_para('• 4.4.3 攻击链构建：邮件路径追踪、URL重定向链、IOC关联分析', True)
    
    add_title('4.5 前端界面实现', 2)
    add_para('• 4.5.1 技术选型：Bootstrap 5响应式框架、Chart.js数据可视化、Inter字体（Apple风格）', True)
    add_para('• 4.5.2 页面设计：数据大屏、控制台、配置管理、对抗测试', True)
    
    # 第五章
    add_title('第五章 系统测试与结果分析', 1)
    
    add_title('5.1 实验环境搭建', 2)
    add_title('5.1.1 硬件环境', 3)
    add_para('• CPU: Intel Core i5-10400', True)
    add_para('• 内存: 16GB DDR4', True)
    add_para('• 存储: 256GB SSD', True)
    add_para('• 模拟中小企业普通服务器配置', True)
    
    add_title('5.1.2 软件环境', 3)
    add_para('• 操作系统: Windows 10 / Ubuntu 20.04', True)
    add_para('• Python: 3.8+', True)
    add_para('• 依赖库: Flask 2.x, LightGBM 3.x, whois, requests', True)
    
    add_title('5.2 数据集与评估指标', 2)
    add_title('5.2.1 数据集', 3)
    add_para('• 正常邮件: Enron数据集（5000封）', True)
    add_para('• 垃圾邮件: SpamAssassin数据集（3000封）', True)
    add_para('• 钓鱼邮件: 自建数据集（2000封）', True)
    
    add_title('5.2.2 评估指标', 3)
    add_para('• 准确率（Accuracy）、精确率（Precision）、召回率（Recall）、F1值（F1-Score）', True)
    add_para('• 检测耗时、内存占用', True)
    
    add_title('5.3 实验结果与分析', 2)
    add_title('5.3.1 检测效果对比', 3)
    add_para('方法对比：', True)
    add_para('• 纯规则方法：准确率85.2%，F1值80.7%', True)
    add_para('• 纯模型方法：准确率92.5%，F1值90.5%', True)
    add_para('• 本文方法：准确率98.5%，F1值97.0%', True)
    
    add_title('5.3.2 资源消耗测试', 3)
    add_para('本文系统 vs BERT方案：', True)
    add_para('• 模型大小：<1MB vs >400MB', True)
    add_para('• 内存占用：~200MB vs >2GB', True)
    add_para('• 单封检测耗时：<2s vs >5s', True)
    add_para('• GPU需求：无 vs 推荐', True)
    
    add_title('5.3.3 对抗样本测试', 3)
    add_para('• 同形字替换攻击：检测率92%', True)
    add_para('• 零宽字符攻击：检测率95%', True)
    add_para('• 同义词替换攻击：检测率88%', True)
    add_para('• 组合攻击：检测率85%', True)
    
    add_title('5.3.4 典型案例分析', 3)
    add_para('• 案例一：银行钓鱼邮件检测全过程', True)
    add_para('• 案例二：身份冒充邮件的溯源追踪', True)
    add_para('• 案例三：边界样本的处理过程', True)
    
    # 第六章
    add_title('第六章 总结与展望', 1)
    
    add_title('6.1 论文工作总结', 2)
    add_para('• 实现了轻量化钓鱼邮件检测系统，适合中小企业部署', True)
    add_para('• 提出了三维度融合检测架构，准确率达到98.5%', True)
    add_para('• 实现了完整的溯源追踪功能，增强取证能力', True)
    add_para('• 系统资源占用低，无需专业运维', True)
    
    add_title('6.2 系统不足之处', 2)
    add_para('• 依赖第三方API（VirusTotal）的稳定性', True)
    add_para('• 规则库需要定期更新维护', True)
    add_para('• 图片类钓鱼邮件检测能力有限', True)
    add_para('• 暂不支持分布式部署', True)
    
    add_title('6.3 未来展望', 2)
    add_para('• 引入多模态模型检测图片钓鱼', True)
    add_para('• 增加自适应学习机制，减少人工干预', True)
    add_para('• 支持更多邮件协议（POP3、Exchange）', True)
    add_para('• 构建云端威胁情报共享平台', True)
    
    # 参考文献
    add_title('参考文献', 1)
    add_para('[1] APWG. Phishing Activity Trends Report[R]. 2023.', True)
    add_para('[2] Verizon. Data Breach Investigations Report[R]. 2023.', True)
    add_para('[3] Ke G, Meng Q, Finley T, et al. LightGBM: A Highly Efficient Gradient Boosting Decision Tree[C]. NIPS, 2017.', True)
    add_para('[4] SpamAssassin. The Apache SpamAssassin Project[EB/OL]. https://spamassassin.apache.org/', True)
    add_para('[5] 邮件安全协议RFC标准: SPF(RFC7208), DKIM(RFC6376), DMARC(RFC7489)', True)
    
    # 致谢
    add_title('致谢', 1)
    
    # 附录
    add_title('附录', 1)
    add_para('• 附录A：系统部署指南', True)
    add_para('• 附录B：API接口文档', True)
    add_para('• 附录C：核心代码清单', True)
    
    # 保存文档
    doc.save('docs/论文大纲.docx')
    print('论文大纲已导出为 docs/论文大纲.docx')

if __name__ == '__main__':
    create_thesis_outline()
