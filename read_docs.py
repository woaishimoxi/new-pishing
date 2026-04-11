# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
import os

ref_dir = r'f:/student/毕设/面向中小型企业的轻量化钓鱼邮件检测与溯源系统设计与实现 - 彻底大改 - trae/参考'
files = [f for f in os.listdir(ref_dir) if f.endswith('.docx')]

for fname in files:
    print('='*70)
    print(f'文件: {fname}')
    print('='*70)
    doc = Document(os.path.join(ref_dir, fname))
    
    # 读取段落
    for para in doc.paragraphs[:150]:
        if para.text.strip():
            print(para.text)
    
    # 读取表格
    for i, table in enumerate(doc.tables):
        print(f'\n--- 表格 {i+1} ---')
        for row in table.rows:
            try:
                row_text = [cell.text.strip() for cell in row.cells]
                if any(row_text):
                    print(' | '.join(row_text))
            except Exception:
                pass
    print()


